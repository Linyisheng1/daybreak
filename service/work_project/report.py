"""Report generation service — produces a professional Word (.docx) security assessment report.

Receives a snapshot dict (from WorkProjectRecordSnapshotSchema.model_dump()),
returns (docx_bytes, filename).
"""
from __future__ import annotations

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Chinese label mappings ──────────────────────────────────────────────────

PROJECT_TYPE = {"penetration_test": "渗透测试", "source_code_audit": "代码审计"}
PROJECT_STATUS = {"working": "进行中", "completed": "已完成", "canceled": "已取消"}

TASK_STATUS = {"todo": "待办", "in_progress": "进行中", "blocked": "受阻", "done": "完成"}

ASSET_TYPE = {"service": "服务", "domain": "域名", "network": "网络", "binary": "二进制"}
ASSET_ORIGIN = {"scope": "范围", "discovered": "发现"}

SEVERITY_LABEL = {
    "critical": "严重", "high": "高危", "medium": "中危",
    "low": "低危", "info": "信息",
}
SEVERITY_ORDER = ["critical", "high", "medium", "low", "info"]
SEVERITY_BG = {
    "critical": "FDE8E8",
    "high":    "FEEBC8",
    "medium":  "FEFCBF",
    "low":     "BEE3F8",
    "info":    "EDF2F7",
}
SEVERITY_TEXT_COLOR = {
    "critical": RGBColor(0x9B, 0x1C, 0x1C),
    "high":    RGBColor(0x9C, 0x42, 0x22),
    "medium":  RGBColor(0x97, 0x5A, 0x10),
    "low":     RGBColor(0x2B, 0x6C, 0xB0),
    "info":    RGBColor(0x4A, 0x55, 0x68),
}

FINDING_STATUS = {"suspected": "疑似", "validated": "已确认", "false_positive": "误报"}
FINDING_STATUS_BG = {"suspected": "FEFCBF", "validated": "C6F6D5", "false_positive": "EDF2F7"}

EDGE_TYPE = {
    "related": "关联", "resolves_to": "解析至", "hosts": "托管",
    "connects_to": "连接至", "trusts": "信任",
    "exploits": "利用", "pivots_to": "跳转至", "leads_to": "导向",
}

ATTACK_PATH_STATUS = {
    "suspected": "疑似", "validated": "已确认",
    "blocked": "已阻止", "closed": "已关闭",
}


# ── Helper utilities ────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_run(paragraph, text: str, size: float = 10.5, bold: bool = False,
             color: RGBColor | None = None, font_name: str = "SimSun",
             italic: bool = False):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = color
    return run


def _add_heading_styled(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "SimHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
    return heading


def _add_info_row(table, row_idx: int, label: str, value: str,
                  label_width: Cm | None = None):
    cell_l = table.rows[row_idx].cells[0]
    cell_r = table.rows[row_idx].cells[1]
    cell_l.text = ""
    cell_r.text = ""
    p_l = cell_l.paragraphs[0]
    _add_run(p_l, label, size=10.5, bold=True, font_name="SimHei")
    _set_cell_shading(cell_l, "EDF2F7")
    if label_width:
        cell_l.width = label_width
    p_r = cell_r.paragraphs[0]
    _add_run(p_r, value, size=10.5)


def _add_styled_table(doc: Document, headers: list[str], rows: list[list[str]],
                      col_widths: list[float] | None = None,
                      severity_col: int | None = None,
                      severity_values: list[str] | None = None):
    if not rows:
        p = doc.add_paragraph("（无数据）")
        _add_run(p, "（无数据）", size=10.5)
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, h, size=10, bold=True, font_name="SimHei", color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_shading(cell, "2D3748")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_run(p, str(val) if val is not None else "", size=9)
            if severity_col is not None and c_idx == severity_col and severity_values:
                sev_key = severity_values[r_idx] if r_idx < len(severity_values) else None
                if sev_key and sev_key in SEVERITY_BG:
                    _set_cell_shading(cell, SEVERITY_BG[sev_key])

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def _add_separator(doc: Document):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CBD5E0"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ── Section builders ────────────────────────────────────────────────────────

def _build_cover(doc: Document, project: dict):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "安全评估报告", size=28, bold=True, font_name="SimHei",
             color=RGBColor(0x1a, 0x36, 0x5d))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, project.get("name", "未命名项目"), size=20, bold=True,
             font_name="SimHei", color=RGBColor(0x2d, 0x37, 0x48))

    doc.add_paragraph()

    owners_list = project.get("owners", [])
    if owners_list and isinstance(owners_list[0], dict):
        owner_names = ", ".join(o.get("username", str(o)) for o in owners_list)
    else:
        owner_names = ", ".join(str(o) for o in owners_list)

    info_lines = [
        f"项目类型：{PROJECT_TYPE.get(project.get('type', ''), project.get('type', '—'))}",
        f"项目状态：{PROJECT_STATUS.get(project.get('status', ''), project.get('status', '—'))}",
        f"负责人：{owner_names or '—'}",
        f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, line, size=13, color=RGBColor(0x4a, 0x55, 0x68))

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "破晓 安全评估平台", size=14, bold=True, font_name="SimHei",
             color=RGBColor(0x4a, 0x55, 0x68))

    doc.add_page_break()


def _build_toc(doc: Document, findings: list):
    _add_heading_styled(doc, "目录", level=1)
    doc.add_paragraph()

    # Count by severity
    sev_counts: dict[str, int] = {}
    for f in findings:
        sev = f.get("severity", "info")
        sev_counts[sev] = sev_counts.get(sev, 0) + 1

    toc_items = [
        "一、项目概述",
        "二、漏洞统计概览",
        "三、漏洞详情",
    ]

    for sev_key in SEVERITY_ORDER:
        count = sev_counts.get(sev_key, 0)
        if count:
            sev_label = SEVERITY_LABEL.get(sev_key, sev_key)
            toc_items.append(f"    {sev_label}漏洞（{count} 项）")

    toc_items.extend([
        "四、资产清单",
        "五、攻击路径",
        "六、修复建议总结",
    ])

    for item in toc_items:
        p = doc.add_paragraph()
        _add_run(p, item, size=12, font_name="SimHei",
                 color=RGBColor(0x2d, 0x37, 0x48))

    doc.add_page_break()


def _build_overview(doc: Document, project: dict, records: dict):
    _add_heading_styled(doc, "一、项目概述", level=1)

    assets = records.get("assets", [])
    findings = records.get("findings", [])
    graph = records.get("graph", {})
    edges = graph.get("edges", [])
    attack_paths = graph.get("attack_paths", [])

    owners_list = project.get("owners", [])
    if owners_list and isinstance(owners_list[0], dict):
        owner_names = ", ".join(o.get("username", str(o)) for o in owners_list)
    else:
        owner_names = ", ".join(str(o) for o in owners_list)

    overview_items = [
        ("项目名称", project.get("name", "—")),
        ("项目类型", PROJECT_TYPE.get(project.get("type", ""), project.get("type", "—"))),
        ("项目状态", PROJECT_STATUS.get(project.get("status", ""), project.get("status", "—"))),
        ("项目描述", project.get("description") or "—"),
        ("负责人", owner_names or "—"),
        ("整体进度", f"{project.get('progress', 0)}%"),
        ("资产数量", str(len(assets))),
        ("漏洞数量", str(len(findings))),
        ("攻击路径数量", str(len(attack_paths))),
    ]

    table = doc.add_table(rows=len(overview_items), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(overview_items):
        _add_info_row(table, i, label, value, label_width=Cm(3.5))

    doc.add_paragraph()


def _build_severity_summary(doc: Document, findings: list):
    _add_heading_styled(doc, "二、漏洞统计概览", level=1)

    sev_counts: dict[str, int] = {}
    for f in findings:
        sev = f.get("severity", "info")
        sev_counts[sev] = sev_counts.get(sev, 0) + 1

    status_counts: dict[str, int] = {}
    for f in findings:
        st = f.get("status", "suspected")
        status_counts[st] = status_counts.get(st, 0) + 1

    # Severity summary table
    _add_heading_styled(doc, "漏洞严重程度分布", level=2)
    sev_rows = []
    sev_vals = []
    for sev_key in SEVERITY_ORDER:
        count = sev_counts.get(sev_key, 0)
        sev_rows.append([SEVERITY_LABEL.get(sev_key, sev_key), str(count),
                         "●" if count > 0 else "○"])
        sev_vals.append(sev_key)

    table = _add_styled_table(doc, ["严重程度", "数量", "状态"], sev_rows,
                              col_widths=[4, 3, 3],
                              severity_col=0, severity_values=sev_vals)

    doc.add_paragraph()

    # Status summary table
    _add_heading_styled(doc, "漏洞验证状态分布", level=2)
    status_rows = []
    for st_key in ["validated", "suspected", "false_positive"]:
        count = status_counts.get(st_key, 0)
        status_rows.append([FINDING_STATUS.get(st_key, st_key), str(count)])
    _add_styled_table(doc, ["验证状态", "数量"], status_rows, col_widths=[5, 3])

    doc.add_paragraph()

    # Risk assessment summary
    critical = sev_counts.get("critical", 0)
    high = sev_counts.get("high", 0)
    if critical > 0:
        risk_level = "极高"
        risk_desc = f"存在 {critical} 个严重漏洞，需立即修复"
    elif high > 0:
        risk_level = "高"
        risk_desc = f"存在 {high} 个高危漏洞，建议尽快修复"
    elif sev_counts.get("medium", 0) > 0:
        risk_level = "中"
        risk_desc = "存在中危漏洞，建议计划修复"
    else:
        risk_level = "低"
        risk_desc = "风险较低，建议关注信息类发现"

    p = doc.add_paragraph()
    _add_run(p, "整体风险评估：", size=12, bold=True, font_name="SimHei")
    risk_color = RGBColor(0x9B, 0x1C, 0x1C) if critical or high else RGBColor(0x97, 0x5A, 0x10)
    _add_run(p, f"{risk_level}风险", size=12, bold=True, font_name="SimHei", color=risk_color)
    _add_run(p, f" — {risk_desc}", size=11)

    doc.add_page_break()


def _build_finding_detail(doc: Document, finding: dict, idx: int, assets: list):
    """Build a detailed finding section with POC and reproduction steps."""
    asset_map = {a.get("id"): a for a in assets}

    title = finding.get("title", "未命名漏洞")
    sev = finding.get("severity", "info")
    sev_label = SEVERITY_LABEL.get(sev, sev)
    status = finding.get("status", "suspected")
    status_label = FINDING_STATUS.get(status, status)
    description = finding.get("description", "")
    impact = finding.get("impact", "")
    asset_id = finding.get("asset_id")
    agent_code = finding.get("created_by_agent_code", "")
    created_at = finding.get("created_at", "")

    # Heading with severity badge
    _add_heading_styled(doc, f"漏洞 {idx}：{title}", level=2)

    # Severity & status badge table
    badge_table = doc.add_table(rows=1, cols=4)
    badge_table.style = "Table Grid"

    # Severity cell
    sev_cell = badge_table.rows[0].cells[0]
    sev_cell.text = ""
    p = sev_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f" {sev_label} ", size=11, bold=True, font_name="SimHei",
             color=SEVERITY_TEXT_COLOR.get(sev, RGBColor(0x4A, 0x55, 0x68)))
    if sev in SEVERITY_BG:
        _set_cell_shading(sev_cell, SEVERITY_BG[sev])
    sev_cell.width = Cm(2.5)

    # Status cell
    st_cell = badge_table.rows[0].cells[1]
    st_cell.text = ""
    p = st_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f" {status_label} ", size=11, bold=True, font_name="SimHei")
    if status in FINDING_STATUS_BG:
        _set_cell_shading(st_cell, FINDING_STATUS_BG[status])
    st_cell.width = Cm(2.5)

    # Agent code
    agent_cell = badge_table.rows[0].cells[2]
    agent_cell.text = ""
    p = agent_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"发现者：{agent_code or '—'}", size=10)
    agent_cell.width = Cm(3.5)

    # Date
    date_cell = badge_table.rows[0].cells[3]
    date_cell.text = ""
    p = date_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = created_at[:10] if created_at else "—"
    _add_run(p, f"发现时间：{date_str}", size=10)
    date_cell.width = Cm(4.5)

    doc.add_paragraph()

    # Affected asset
    if asset_id:
        asset = asset_map.get(asset_id, {})
        asset_name = asset.get("identifier", str(asset_id))
        asset_type = ASSET_TYPE.get(asset.get("type", ""), asset.get("type", ""))
        host = asset.get("host", "")
        port = asset.get("port", "")

        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = "Table Grid"
        _add_info_row(info_table, 0, "受影响资产", f"{asset_name}（{asset_type}）{host}:{port}",
                      label_width=Cm(3))
        doc.add_paragraph()

    # ── 漏洞描述 (full, not truncated) ──
    p = doc.add_paragraph()
    _add_run(p, "漏洞描述", size=11, bold=True, font_name="SimHei",
             color=RGBColor(0x1a, 0x36, 0x5d))

    if description:
        # Split by newlines for readability
        for line_idx, line in enumerate(description.split("\n")):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            # Check if it's a bullet/list item
            if line.startswith("- ") or line.startswith("• "):
                _add_run(p, "  ", size=10.5)
                _add_run(p, line, size=10.5)
                p.paragraph_format.left_indent = Cm(0.5)
            elif line.startswith(tuple("0123456789")) and "." in line[:4]:
                _add_run(p, "  ", size=10.5)
                _add_run(p, line, size=10.5)
                p.paragraph_format.left_indent = Cm(0.5)
            else:
                _add_run(p, line, size=10.5)
    else:
        p = doc.add_paragraph()
        _add_run(p, "（无描述）", size=10.5, color=RGBColor(0xA0, 0xAE, 0xC0))

    doc.add_paragraph()

    # ── 影响分析 (full, not truncated) ──
    if impact:
        p = doc.add_paragraph()
        _add_run(p, "影响分析", size=11, bold=True, font_name="SimHei",
                 color=RGBColor(0x9B, 0x1C, 0x1C))

        for line in impact.split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            if line.startswith(tuple("0123456789")) and ("." in line[:4] or ")" in line[:4]):
                _add_run(p, "  ", size=10.5)
                _add_run(p, line, size=10.5)
                p.paragraph_format.left_indent = Cm(0.5)
            else:
                _add_run(p, line, size=10.5)

        doc.add_paragraph()

    # ── POC / 复现步骤 (extracted from description) ──
    p = doc.add_paragraph()
    _add_run(p, "POC / 复现步骤", size=11, bold=True, font_name="SimHei",
             color=RGBColor(0x2B, 0x6C, 0xB0))

    # Extract POC from description - look for URLs, paths, HTTP methods
    poc_lines = []
    url_lines = []
    method_lines = []

    all_text = f"{description}\n{impact}"

    for line in all_text.split("\n"):
        line_stripped = line.strip()
        if not line_stripped:
            continue
        # Detect HTTP URLs
        if "http://" in line_stripped or "https://" in line_stripped:
            url_lines.append(line_stripped)
        # Detect HTTP methods
        if line_stripped.upper().startswith(("GET ", "POST ", "PUT ", "DELETE ", "PATCH ", "OPTIONS ")):
            method_lines.append(line_stripped)
        # Detect path patterns
        if line_stripped.startswith("/") and not line_stripped.startswith("//"):
            url_lines.append(line_stripped)
        # Detect specific exploit indicators
        if any(kw in line_stripped.lower() for kw in ["curl", "wget", "payload", "exploit", "poc", "request", "response"]):
            poc_lines.append(line_stripped)

    if url_lines or method_lines or poc_lines:
        # Build a POC box
        poc_table = doc.add_table(rows=1, cols=1)
        poc_table.style = "Table Grid"
        poc_cell = poc_table.rows[0].cells[0]
        poc_cell.text = ""
        _set_cell_shading(poc_cell, "F7FAFC")

        first = True
        for url in url_lines:
            if not first:
                p = poc_cell.add_paragraph()
            else:
                p = poc_cell.paragraphs[0]
                first = False
            _add_run(p, "  目标地址：", size=10, bold=True, font_name="SimHei")
            _add_run(p, url, size=10, color=RGBColor(0x2B, 0x6C, 0xB0))

        for method in method_lines:
            p = poc_cell.add_paragraph()
            _add_run(p, "  请求方式：", size=10, bold=True, font_name="SimHei")
            _add_run(p, method, size=10, color=RGBColor(0x2B, 0x6C, 0xB0))

        for poc in poc_lines:
            p = poc_cell.add_paragraph()
            _add_run(p, "  ", size=10)
            _add_run(p, poc, size=10, color=RGBColor(0x97, 0x5A, 0x10))

        # Add reproduction steps
        if url_lines:
            p = poc_cell.add_paragraph()
            p = poc_cell.add_paragraph()
            _add_run(p, "  复现步骤：", size=10, bold=True, font_name="SimHei")

            step_num = 1
            # Step 1: access the URL
            for url in url_lines[:3]:
                p = poc_cell.add_paragraph()
                _add_run(p, f"    {step_num}. 访问 {url}", size=10)
                step_num += 1

            # Step 2: observe result
            if description:
                p = poc_cell.add_paragraph()
                _add_run(p, f"    {step_num}. 观察返回数据，确认漏洞存在", size=10)
                step_num += 1

            # Step 3: impact
            if impact:
                p = poc_cell.add_paragraph()
                _add_run(p, f"    {step_num}. 验证影响范围", size=10)
    else:
        # Generic POC section from description
        p = doc.add_paragraph()
        _add_run(p, "  根据漏洞描述，可按以下方式验证：", size=10.5, color=RGBColor(0x4A, 0x55, 0x68))
        step_num = 1
        for line in description.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("• "):
                p = doc.add_paragraph()
                _add_run(p, f"    {step_num}. {line[2:]}", size=10.5)
                step_num += 1
            elif line.startswith(tuple("0123456789")) and "." in line[:4]:
                p = doc.add_paragraph()
                _add_run(p, f"    {step_num}. {line}", size=10.5)
                step_num += 1

    doc.add_paragraph()

    # ── 修复建议 ──
    p = doc.add_paragraph()
    _add_run(p, "修复建议", size=11, bold=True, font_name="SimHei",
             color=RGBColor(0x27, 0x67, 0x49))

    # Generate remediation based on severity and description
    remediation = _generate_remediation(finding)
    for line in remediation.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        _add_run(p, "  ✓ ", size=10.5, bold=True, color=RGBColor(0x27, 0x67, 0x49))
        _add_run(p, line, size=10.5)

    doc.add_paragraph()
    _add_separator(doc)


def _generate_remediation(finding: dict) -> str:
    """Generate targeted remediation advice based on finding title and content."""
    desc = (finding.get("description", "") + " " + finding.get("impact", "")).lower()
    title = finding.get("title", "").lower()
    lines = []

    # Priority-ordered classification — first strong match wins the primary advice,
    # then we add a few cross-cutting items if they're genuinely relevant.

    # 1. Padding Oracle / crypto attack (very specific — must be first)
    if any(kw in title for kw in ["padding oracle", "padding", "oracle"]) or \
       any(kw in desc for kw in ["padding oracle", "padding error", "vaudenay"]):
        lines.append("在解密前必须先验证 MAC/签名（Encrypt-then-MAC 或 AEAD 模式），绝不允许仅依赖解密结果判断")
        lines.append("使用标准的 ASP.NET DataProtection 机制替代自定义加密实现")
        lines.append("确保所有加密操作使用随机 IV/Nonce，避免确定性加密")
        lines.append("统一错误响应，使攻击者无法区分填充错误与内容错误")

    # 2. Replay attack
    elif any(kw in title for kw in ["重放", "replay"]):
        lines.append("在认证请求中引入 nonce（一次性随机数）和时间戳验证")
        lines.append("服务端记录已使用的 nonce，拒绝重复请求")
        lines.append("设置请求有效期（如 5 分钟），超时自动失效")
        lines.append("实施请求签名机制，防止参数篡改")

    # 3. Error stack trace / info disclosure (must come before auth/crypto to avoid mis-match)
    elif any(kw in title for kw in ["错误堆栈", "堆栈泄露", "源代码路径", "架构信息"]):
        lines.append("配置自定义错误页面，禁止在生产环境返回详细堆栈跟踪")
        lines.append("实施全局异常处理，统一返回通用错误响应")
        lines.append("移除或隐藏不必要的版本信息和内部架构细节")
        lines.append("实施安全头部（Security Headers）配置")

    # 4. Log/audit file exposure
    elif any(kw in title for kw in ["日志", "审计日志", "log expos", "audit log", "明文暴露"]):
        lines.append("将日志和审计文件移至 Web 根目录之外，禁止通过 HTTP 访问")
        lines.append("配置访问控制策略，限制日志目录的访问权限")
        lines.append("对敏感日志文件实施加密存储和传输")
        lines.append("避免在日志中记录完整的认证请求/响应和敏感数据")

    # 5. Auth token/credential exposure or weak crypto
    elif any(kw in title for kw in ["授权码", "认证数据加密", "静态iv", "确定性加密", "自定义加密", "加密协议", "链式授权", "不可撤销", "认证加密", "密钥"]):
        lines.append("使用标准加密库和协议（如 AES-GCM），避免自定义加密实现")
        lines.append("每次加密使用随机 IV/Nonce，确保密文不可预测")
        lines.append("实施 MAC 验证机制（Encrypt-then-MAC 或 AEAD），在解密前验证完整性")
        lines.append("为授权码设置过期时间和一次性使用限制")
        lines.append("建立授权码撤销机制，支持紧急吊销已泄露的授权码")

    # 4. Rate limiting
    elif any(kw in title for kw in ["速率限制", "rate limit", "无限制"]) or \
         any(kw in desc for kw in ["速率限制", "rate limit", "无限次"]):
        lines.append("实施 API 速率限制（Rate Limiting），根据端点敏感度设置不同阈值")
        lines.append("对认证接口设置严格的调用频率限制（如每分钟 5 次）")
        lines.append("记录并监控异常高频请求，触发告警机制")
        lines.append("考虑实施 IP 黑名单/灰名单机制")

    # 5. Directory listing
    elif any(kw in title for kw in ["目录列表", "目录浏览", "directory brows", "directory list"]):
        lines.append("关闭 Web 服务器的目录列表功能（Directory Browsing/Autoindex）")
        lines.append("在 IIS 中禁用 Directory Browsing，或在 Nginx/Apache 中移除 autoindex 配置")
        lines.append("确保所有目录都有默认首页文件（index.html 等）")

    # 6. SQL / database
    elif any(kw in title for kw in ["sql", "数据库", "database"]):
        lines.append("使用参数化查询防止 SQL 注入")
        lines.append("限制数据库用户权限，实施最小权限原则")
        lines.append("禁止在日志中记录完整 SQL 查询语句")

    # 7. Security headers missing
    elif any(kw in title for kw in ["安全响应头", "security header", "响应头缺失"]):
        lines.append("配置完整的安全响应头：X-Content-Type-Options, X-Frame-Options, CSP, HSTS 等")
        lines.append("启用 Strict-Transport-Security（HSTS）强制 HTTPS")
        lines.append("配置 Content-Security-Policy（CSP）防止 XSS 攻击")
        lines.append("设置 X-Content-Type-Options: nosniff 防止 MIME 类型嗅探")

    # 8. Debug tool / profiler
    elif any(kw in title for kw in ["miniprofiler", "profiler", "调试工具"]):
        lines.append("在生产环境中禁用所有调试工具和性能分析器")
        lines.append("移除 MiniProfiler 相关的 HTTP 响应头")
        lines.append("使用环境变量或配置开关控制调试功能的启用/禁用")

    # 9. Prometheus / metrics endpoint
    elif any(kw in title for kw in ["metrics", "prometheus", "指标"]):
        lines.append("为 /metrics 端点添加认证和访问控制")
        lines.append("限制 Prometheus 端点仅对监控网络开放")
        lines.append("过滤敏感指标数据，避免暴露内部架构信息")

    # 10. License / endpoint exposure
    elif any(kw in title for kw in ["license", "端点暴露", "接口可访问"]):
        lines.append("对管理接口和内部端点实施网络隔离和访问控制")
        lines.append("添加认证机制保护所有非公开端点")
        lines.append("实施 IP 白名单限制管理接口的访问来源")

    # 11. DataProtection / key management
    elif any(kw in title for kw in ["dataprotection", "密钥未加密", "key management"]):
        lines.append("使用标准 ASP.NET DataProtection API，避免自定义实现")
        lines.append("启用密钥加密存储（使用 Windows DPAPI 或证书保护密钥）")
        lines.append("定期轮换加密密钥，设置合理的密钥过期策略")
        lines.append("确保密钥仅在受保护的环境中存储和访问")

    # 12. Generic info disclosure fallback
    elif any(kw in desc for kw in ["信息泄露", "暴露", "exposure", "disclosure", "泄露"]):
        lines.append("移除或隐藏不必要的版本信息和内部架构细节")
        lines.append("配置自定义错误页面，避免暴露堆栈跟踪")
        lines.append("实施安全头部（Security Headers）配置")

    # 15. Complete fallback
    else:
        lines.append("对受影响的系统组件进行安全加固")
        lines.append("限制不必要的网络访问和服务暴露")
        lines.append("定期进行安全评估和漏洞扫描")

    # Add severity-based urgency
    sev = finding.get("severity", "info")
    if sev == "critical":
        lines.append("【紧急】此漏洞风险极高，建议在 24 小时内修复")
    elif sev == "high":
        lines.append("【重要】此漏洞风险较高，建议在 72 小时内修复")
    elif sev == "medium":
        lines.append("【建议】此漏洞建议在 1 周内修复")

    return "\n".join(lines)


def _build_findings_detailed(doc: Document, records: dict):
    """Section 3: Detailed findings with POC and reproduction steps."""
    _add_heading_styled(doc, "三、漏洞详情", level=1)

    findings = records.get("findings", [])
    assets = records.get("assets", [])

    if not findings:
        p = doc.add_paragraph()
        _add_run(p, "本次评估未发现安全漏洞。", size=12, font_name="SimHei",
                 color=RGBColor(0x27, 0x67, 0x49))
        return

    # Group by severity
    grouped: dict[str, list] = {}
    for f in findings:
        sev = f.get("severity", "info")
        grouped.setdefault(sev, []).append(f)

    global_idx = 1
    for sev_key in SEVERITY_ORDER:
        group = grouped.get(sev_key, [])
        if not group:
            continue
        sev_label = SEVERITY_LABEL.get(sev_key, sev_key)
        _add_heading_styled(doc, f"{sev_label}漏洞（{len(group)} 项）", level=2)

        for finding in group:
            _build_finding_detail(doc, finding, global_idx, assets)
            global_idx += 1


def _build_assets(doc: Document, records: dict):
    _add_heading_styled(doc, "四、资产清单", level=1)

    assets = records.get("assets", [])
    if not assets:
        p = doc.add_paragraph()
        _add_run(p, "（无资产）", size=10.5)
        return

    headers = ["序号", "资产标识", "类型", "来源", "主机", "端口", "横幅"]
    rows = []
    for i, a in enumerate(assets, 1):
        identifier = a.get("identifier", "—")
        atype = ASSET_TYPE.get(a.get("type", ""), a.get("type", "—"))
        origin = ASSET_ORIGIN.get(a.get("origin", ""), a.get("origin", "—"))
        host = a.get("host") or "—"
        port = str(a.get("port")) if a.get("port") else "—"
        extra = a.get("extra", {})
        banner = a.get("banner") or (extra.get("banner") if isinstance(extra, dict) else "") or "—"
        if len(banner) > 80:
            banner = banner[:80] + "…"
        rows.append([str(i), identifier, atype, origin, host, port, banner])

    _add_styled_table(doc, headers, rows, col_widths=[1.2, 4, 1.5, 1.5, 3, 1.5, 4])


def _build_attack_paths(doc: Document, records: dict):
    _add_heading_styled(doc, "五、攻击路径", level=1)

    graph = records.get("graph", {})
    attack_paths = graph.get("attack_paths", [])
    edges = graph.get("edges", [])
    attack_path_steps = graph.get("attack_path_steps", [])

    if not attack_paths:
        p = doc.add_paragraph()
        _add_run(p, "（无攻击路径）", size=10.5)
        return

    edge_map = {e.get("id"): e for e in edges}
    assets = records.get("assets", [])
    asset_map = {a.get("id"): a for a in assets}

    steps_by_path: dict[int, list] = {}
    for step in attack_path_steps:
        pid = step.get("path_id")
        steps_by_path.setdefault(pid, []).append(step)
    for pid in steps_by_path:
        steps_by_path[pid].sort(key=lambda s: s.get("sequence", 0))

    for i, path in enumerate(attack_paths, 1):
        title = path.get("title") or f"攻击路径 {i}"
        status = ATTACK_PATH_STATUS.get(path.get("status", ""), path.get("status", "—"))
        summary = path.get("summary") or ""

        _add_heading_styled(doc, f"路径 {i}：{title}（{status}）", level=2)

        if summary:
            p = doc.add_paragraph()
            _add_run(p, f"摘要：{summary}", size=10.5)

        path_id = path.get("id")
        path_steps = steps_by_path.get(path_id, [])
        if path_steps:
            headers = ["步骤", "源→目标", "关系"]
            rows = []
            for s_idx, step in enumerate(path_steps, 1):
                edge = edge_map.get(step.get("edge_id"), {})
                src_id = edge.get("source_asset_id")
                tgt_id = edge.get("target_asset_id")
                src_name = asset_map.get(src_id, {}).get("identifier", str(src_id or "?"))
                tgt_name = asset_map.get(tgt_id, {}).get("identifier", str(tgt_id or "?"))
                if len(src_name) > 40:
                    src_name = src_name[:40] + "…"
                if len(tgt_name) > 40:
                    tgt_name = tgt_name[:40] + "…"
                rel = EDGE_TYPE.get(edge.get("type", ""), edge.get("type", "—"))
                rows.append([str(s_idx), f"{src_name} → {tgt_name}", rel])
            _add_styled_table(doc, headers, rows, col_widths=[1.5, 8, 3])

        doc.add_paragraph()


def _build_remediation_summary(doc: Document, findings: list):
    """Section 6: Consolidated remediation summary."""
    _add_heading_styled(doc, "六、修复建议总结", level=1)

    if not findings:
        p = doc.add_paragraph()
        _add_run(p, "无需修复。", size=12, font_name="SimHei", color=RGBColor(0x27, 0x67, 0x49))
        return

    # Priority-ordered remediation
    sev_counts: dict[str, int] = {}
    for f in findings:
        sev = f.get("severity", "info")
        sev_counts[sev] = sev_counts.get(sev, 0) + 1

    # Urgency table
    urgency_rows = []
    for sev_key in SEVERITY_ORDER:
        count = sev_counts.get(sev_key, 0)
        if count == 0:
            continue
        sev_label = SEVERITY_LABEL.get(sev_key, sev_key)
        if sev_key == "critical":
            timeline = "24 小时内"
        elif sev_key == "high":
            timeline = "72 小时内"
        elif sev_key == "medium":
            timeline = "1 周内"
        else:
            timeline = "建议处理"
        urgency_rows.append([sev_label, str(count), timeline])

    _add_heading_styled(doc, "修复优先级", level=2)
    _add_styled_table(doc, ["严重程度", "漏洞数量", "建议修复时限"], urgency_rows,
                      col_widths=[4, 3, 4],
                      severity_col=0, severity_values=[k for k in SEVERITY_ORDER if sev_counts.get(k, 0) > 0])

    doc.add_paragraph()

    # General recommendations
    _add_heading_styled(doc, "通用安全建议", level=2)
    general_recs = [
        "对所有发现的安全漏洞按优先级进行修复，优先处理严重和高危漏洞",
        "实施安全开发生命周期（SDL），在开发阶段引入安全审查",
        "定期进行安全评估和渗透测试，及时发现新风险",
        "建立漏洞管理流程，跟踪漏洞修复进度和验证结果",
        "加强安全意识培训，提高开发团队的安全编码能力",
    ]
    for rec in general_recs:
        p = doc.add_paragraph()
        _add_run(p, "  ✓ ", size=10.5, bold=True, color=RGBColor(0x27, 0x67, 0x49))
        _add_run(p, rec, size=10.5)

    doc.add_paragraph()

    # Disclaimer
    p = doc.add_paragraph()
    _add_run(p, "免责声明", size=10, bold=True, font_name="SimHei", color=RGBColor(0x71, 0x81, 0x96))
    p = doc.add_paragraph()
    _add_run(p, "本报告仅针对评估期间所观察到的安全状况，不构成对系统整体安全性的保证。"
             "评估结果基于当前系统配置和测试时间点，系统变更后可能产生新的安全风险。"
             "本报告内容仅供内部参考，未经授权不得对外公开。",
             size=9, color=RGBColor(0x71, 0x81, 0x96), italic=True)


# ── Main entry ──────────────────────────────────────────────────────────────

async def generate_report(snapshot: dict) -> tuple[bytes, str]:
    """Generate a Word (.docx) report from a project snapshot.

    Returns (docx_bytes, filename) where filename is URL-safe Chinese.
    """
    project = snapshot.get("project", {})
    records = snapshot.get("records", {})
    findings = records.get("findings", [])

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build sections
    _build_cover(doc, project)
    _build_toc(doc, findings)
    _build_overview(doc, project, records)
    _build_severity_summary(doc, findings)
    _build_findings_detailed(doc, records)
    _build_assets(doc, records)
    _build_attack_paths(doc, records)
    _build_remediation_summary(doc, findings)

    # Save
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    project_name = project.get("name", "未命名")
    safe_name = project_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Daybreak_安全评估报告_{safe_name}_{timestamp}.docx"

    return docx_bytes, filename
